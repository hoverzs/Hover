"""Igehirdetési vázlat Word (.docx) exportja.

Kiemelve az `app.py`-ból (2026-08-13, célarchitektúra-terv 2. fázis, 1.
lépés) — a Textusműhely önálló "Vázlat"/"Gyors vázlat" kártyájának
megszűnésével ez a képesség egyetlen felhasználói belépési ponton, az
Igehirdetési műhely "Igehirdetési vázlat" szekciójából érhető el
(`sermon_workshop_ui.render_outline_section`).

Önálló modulként azért kellett kiemelni (nem elég lett volna lazy
`from app import build_outline_docx`), mert az `app.py` egy tiszta
Streamlit-szkript — `streamlit run app.py` alatt a betöltött modul neve
`__main__`, nem `app`; egy futásidejű `import app` emiatt egy MÁSODIK,
friss példányban futtatná újra a teljes app.py-t (beleértve a
`st.set_page_config(...)` hívást is, ami másodszor hibát dobna). Ez a modul
Streamlit-független (csak `st.session_state`-et olvas), nem importál sem
`app`-ot, sem `sermon_workshop_ui`-t — nincs körkörös import kockázat.

A tartalom-struktúra változatlan: a vázlat törzsét és a kosár/ének
szövegeket soronkénti, egyszerű Markdown-heurisztikával alakítja Word-be
(UTF-8, Calibri — magyar ékezetek)."""

from __future__ import annotations

import io
import re
from datetime import datetime

import streamlit as st

APP_NAME = "TEXTUS"
APP_VERSION = "3.0"
APP_SUBTITLE = "Homiletikai műhely"
APP_TAGLINE = "A szövegtől a szószékig"


def _docx_strip_md_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text or "")


def _docx_add_inline_runs(paragraph, line: str) -> None:
    """`**félkövér**` és sima szöveg — páros `**` felosztással."""
    t = _docx_strip_md_links(line)
    parts = t.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:
            run.bold = True


def _docx_append_markdown_body(doc, text: str) -> None:
    """Markdown-szerű blokk Word-be: címsorok, listák, idézet, üres sorok."""
    if not (text or "").strip():
        p = doc.add_paragraph(style="Intense Quote")
        p.add_run("_Még nem készült vázlat._")
        return
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            doc.add_heading(hm.group(2).strip(), level=level)
            continue
        if stripped.startswith(">"):
            content = stripped.lstrip(">").strip()
            p = doc.add_paragraph(style="Quote")
            _docx_add_inline_runs(p, content)
            continue
        if re.match(r"^[-*+]\s+", stripped):
            content = re.sub(r"^[-*+]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_inline_runs(p, content)
            continue
        if re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _docx_add_inline_runs(p, content)
            continue
        p = doc.add_paragraph()
        _docx_add_inline_runs(p, stripped)


def build_outline_docx() -> bytes:
    """Összeállítja a vázlatkosár + ének Word dokumentumát (bináris .docx)."""
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(f"Prédikációvázlat — {APP_NAME}", level=1)

    igehely = st.session_state.get("last_igehely", "—")
    alkalom = st.session_state.get("last_alkalom", "—")
    stilus = st.session_state.get("last_stilus", "—")
    outline = st.session_state.get("outline", "").strip()
    basket = st.session_state.get("basket", [])
    songs = st.session_state.get("songs", "").strip()
    now = datetime.now().strftime("%Y. %m. %d. %H:%M")

    p_meta = doc.add_paragraph()
    r_l = p_meta.add_run("Igehely: ")
    r_l.bold = True
    r_v = p_meta.add_run(igehely)
    r_v.bold = True
    try:
        r_v.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except Exception:
        r_v.italic = True

    p_al = doc.add_paragraph()
    p_al.add_run("Alkalom: ").bold = True
    _docx_add_inline_runs(p_al, alkalom)

    p_st = doc.add_paragraph()
    p_st.add_run("Homiletikai stílus: ").bold = True
    _docx_add_inline_runs(p_st, stilus)

    doc.add_paragraph(f"Készült: {now}")
    doc.add_paragraph()

    doc.add_heading("Vázlat", level=2)
    _docx_append_markdown_body(doc, outline)

    if basket:
        doc.add_heading("Vázlatkosár — gondolatok a vázlathoz", level=2)
        for source, item in basket:
            doc.add_heading(source, level=3)
            _docx_append_markdown_body(doc, item)

    if songs:
        doc.add_heading("Liturgiai énekajánlás", level=2)
        _docx_append_markdown_body(doc, songs)

    doc.add_paragraph()
    p_f = doc.add_paragraph()
    r_f = p_f.add_run(f"{APP_NAME} v{APP_VERSION} — {APP_SUBTITLE} · {APP_TAGLINE}")
    r_f.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["build_outline_docx"]
