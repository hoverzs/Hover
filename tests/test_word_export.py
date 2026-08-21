# ruff: noqa: E402
"""Word (.docx) vázlat-export: függőség, tartalom, felhasználói hibaüzenet."""

from __future__ import annotations

import io
import re
import sys
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

WORD_EXPORT_USER_MSG = (
    "A Word-export jelenleg nem érhető el. "
    "Az alkalmazás egyik dokumentumkezelő összetevője hiányzik."
)

SAMPLE_OUTLINE = """\
## Bevezetés
Árvíztűrő tükörfúrógép — magyar ékezetek.

### Fő gondolat
- hiteles tanúságtétel
- közösségi felelősség
1. első lépés
2. második lépés

> „Az Úr az én pásztorom.”

## Lekció
Jn 10,11–16 — A jó pásztor.

## Imádság
Uram, őrizd a nyájat.
"""


@pytest.fixture
def outline_session():
    """Minimális session_state a build_outline_docx számára."""
    return {
        "last_igehely": "Zsolt 23,1–4",
        "last_alkalom": "vasárnapi igehirdetés",
        "last_stilus": "expozíciós",
        "outline": SAMPLE_OUTLINE,
        "basket": [],
        "songs": "",
    }


def test_requirements_pins_python_docx_not_docx_package():
    text = (ROOT / "requirements.txt").read_text(encoding="utf-8")
    assert re.search(r"(?m)^python-docx\s*[>=<~!]=?", text)
    # A PyPI `docx` csomag nem helyettesíti a python-docx-ot
    assert not re.search(r"(?m)^docx\s*[>=<~!]=?", text)


def test_docx_package_importable():
    from docx import Document as Doc

    assert Doc is Document


def test_build_outline_docx_rich_content(outline_session, tmp_path):
    # 2026-08-13: a build_outline_docx az outline_word_export.py-ba került
    # (ld. a modul docstringjét) — a saját `st` referenciáját onnan olvassa.
    import outline_word_export

    class _SS(dict):
        pass

    ss = _SS()
    ss.update(outline_session)

    with patch.object(outline_word_export, "st") as st_mock:
        st_mock.session_state = ss
        data = outline_word_export.build_outline_docx()

    assert isinstance(data, (bytes, bytearray))
    assert len(data) > 1000
    assert data[:2] == b"PK"  # zip/docx

    # Opcionális temp fájl (nem commitoljuk)
    out = tmp_path / "outline_sample.docx"
    out.write_bytes(data)
    assert out.stat().st_size > 1000

    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    assert "Zsolt 23,1–4" in joined or "Zsolt 23" in joined
    assert "Árvíztűrő" in joined
    assert "tükörfúrógép" in joined
    assert "hiteles tanúságtétel" in joined
    assert "első lépés" in joined
    assert "Lekció" in joined or "Jn 10" in joined
    assert "Imádság" in joined or "őrizd" in joined
    assert "pásztorom" in joined

    styles = {p.style.name for p in doc.paragraphs if p.style is not None}
    assert any("Heading" in s or "List" in s for s in styles)


def test_build_outline_docx_empty_optional_sections():
    """Üres kosár / ének: csak a vázlat jelenik meg, nincs crash."""
    import outline_word_export

    class _SS(dict):
        pass

    ss = _SS(
        {
            "last_igehely": "Jn 1,1",
            "last_alkalom": "—",
            "last_stilus": "—",
            "outline": "## Cím\nRövid szöveg.",
            "basket": [],
            "songs": "",
        }
    )

    with patch.object(outline_word_export, "st") as st_mock:
        st_mock.session_state = ss
        data = outline_word_export.build_outline_docx()

    doc = Document(io.BytesIO(data))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "Vázlatkosár" not in joined
    assert "Liturgiai énekajánlás" not in joined
    assert "Rövid szöveg" in joined


def test_build_outline_docx_empty_outline_placeholder():
    import outline_word_export

    class _SS(dict):
        pass

    ss = _SS(
        {
            "last_igehely": "—",
            "last_alkalom": "—",
            "last_stilus": "—",
            "outline": "",
            "basket": [],
            "songs": "",
        }
    )

    with patch.object(outline_word_export, "st") as st_mock:
        st_mock.session_state = ss
        data = outline_word_export.build_outline_docx()

    doc = Document(io.BytesIO(data))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "Még nem készült vázlat" in joined


def test_word_export_ui_message_and_no_install_leak():
    """Felhasználói üzenet + Technikai részletek; nincs pip/path leak.

    2026-08-13 (célarchitektúra-terv, 2. fázis, 1. lépés): a letöltés-gomb
    a Textusműhely önálló "Vázlat" kártyájának megszűnésével átkerült az
    Igehirdetési műhely "Igehirdetési vázlat" szekciójába
    (sermon_workshop_ui.py), a közös `outline_word_export.build_outline_docx`
    hívásával — ugyanazzal a felhasználói szöveggel és hibakezeléssel."""
    src = (ROOT / "sermon_workshop_ui.py").read_text(encoding="utf-8")
    idx = src.find('key="outline_download_docx"')
    assert idx > 0
    block = src[idx : idx + 1200]
    collapsed = re.sub(r"\s+", " ", block)
    assert WORD_EXPORT_USER_MSG in collapsed
    assert "Technikai részletek" in block
    assert "expanded=False" in block
    assert "pip install" not in block
    assert "sys.executable" not in block
    assert "Aktív Python" not in block
    assert "átmenetileg nem érhető el" not in block
