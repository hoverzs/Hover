"""Íróasztal draft DOCX export — formázás, fájlnév, üres állapot, flush."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from writing_desk_data import WRITING_DESK_KEY, set_writing_desk_draft
from writing_desk_docx import (
    build_writing_desk_docx_bytes,
    writing_desk_docx_filename,
)


ROOT = Path(__file__).resolve().parents[1]


def _open(html: str) -> Document:
    blob = build_writing_desk_docx_bytes(html)
    assert blob is not None
    return Document(io.BytesIO(blob))


def _all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def _runs(doc: Document):
    for paragraph in doc.paragraphs:
        yield from paragraph.runs


def test_docx_module_has_no_streamlit():
    src = (ROOT / "writing_desk_docx.py").read_text(encoding="utf-8")
    assert "streamlit" not in src
    assert "from app import" not in src
    ui_src = (ROOT / "writing_desk_ui.py").read_text(encoding="utf-8")
    assert "from docx import" not in ui_src


def test_plain_paragraph():
    doc = _open("<p>1. Belépés</p>")
    assert _all_text(doc).strip().startswith("1. Belépés")


def test_bold_italic_underline():
    doc = _open(
        "<p><strong>félkövér</strong> "
        "<em>dőlt</em> "
        "<u>aláhúzott</u></p>"
    )
    texts = {((r.text or ""), bool(r.bold), bool(r.italic), bool(r.underline)) for r in _runs(doc)}
    assert any(text == "félkövér" and bold for text, bold, _, _ in texts)
    assert any(text == "dőlt" and italic for text, _, italic, _ in texts)
    assert any(text == "aláhúzott" and under for text, _, _, under in texts)


def test_br_inside_paragraph():
    doc = _open("<p>első sor<br>második sor</p>")
    text = _all_text(doc)
    assert "első sor" in text
    assert "második sor" in text
    assert text.count("\n") >= 1


def test_unordered_and_ordered_lists():
    doc = _open(
        "<ul><li>egy</li><li>kettő</li></ul>"
        "<ol><li>első</li><li>második</li></ol>"
    )
    styles = [p.style.name for p in doc.paragraphs if (p.text or "").strip()]
    texts = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
    assert "egy" in texts and "kettő" in texts
    assert "első" in texts and "második" in texts
    assert styles.count("List Bullet") >= 2
    assert styles.count("List Number") >= 2


def test_hungarian_accents():
    doc = _open("<p>árvíztűrő tükörfúrógép</p>")
    assert "árvíztűrő tükörfúrógép" in _all_text(doc)


def test_html_entities_decoded():
    doc = _open("<p>A &amp; B &lt; C &gt; D</p>")
    text = _all_text(doc)
    assert "A & B < C > D" in text
    assert "&amp;" not in text


def test_no_raw_html_in_docx():
    doc = _open("<p>Sima <strong>kiemelt</strong> szöveg</p>")
    blob = build_writing_desk_docx_bytes("<p>Sima <strong>kiemelt</strong> szöveg</p>")
    assert blob is not None
    joined = _all_text(doc)
    assert "<p>" not in joined
    assert "<strong>" not in joined
    assert "</strong>" not in joined
    # A Word XML-ben sincs a felhasználói HTML tag szövegként.
    xml = blob.decode("utf-8", errors="ignore")
    assert "&lt;p&gt;" not in xml
    assert "<strong>" not in xml


def test_paragraph_alignment_exports_to_docx():
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _open(
        '<p style="text-align: center">Közép</p>'
        '<p style="text-align: right">Jobb</p>'
        '<p style="text-align: justify">Sorkizárt</p>'
    )
    texts = {p.text: p.alignment for p in doc.paragraphs if (p.text or "").strip()}
    assert texts["Közép"] == WD_ALIGN_PARAGRAPH.CENTER
    assert texts["Jobb"] == WD_ALIGN_PARAGRAPH.RIGHT
    assert texts["Sorkizárt"] == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert build_writing_desk_docx_bytes("") is None
    assert build_writing_desk_docx_bytes("<p><br></p>") is None
    assert build_writing_desk_docx_bytes("   ") is None


def test_filename_sanitization():
    assert writing_desk_docx_filename("") == "Textus_vazlat.docx"
    assert writing_desk_docx_filename(None) == "Textus_vazlat.docx"
    assert writing_desk_docx_filename("Jn 3,16") == "Textus_Jn_3_16_vazlat.docx"
    assert writing_desk_docx_filename("Jn 3:16") == "Textus_Jn_3_16_vazlat.docx"
    name = writing_desk_docx_filename("Jn 3:16 / extra:cím")
    assert "/" not in name
    assert ":" not in name
    assert name.endswith(".docx")
    assert writing_desk_docx_filename("a" * 200).endswith("_vazlat.docx")
    assert len(writing_desk_docx_filename("a" * 200)) < 80


def test_no_automatic_header_or_reference_in_body():
    doc = _open("<p>1. Belépés</p>")
    text = _all_text(doc)
    assert text.strip() == "1. Belépés"
    assert "Textus" not in text
    assert "Jn" not in text


def test_export_flushes_stale_ccv2_widget_state(monkeypatch):
    import streamlit as st

    import writing_desk_ui
    from writing_desk_ui import WRITING_DESK_DRAFT_WIDGET_KEY

    session: dict = {"last_igehely": "Jn 3,16"}
    set_writing_desk_draft(session, "<p>régi durable</p>")
    session[WRITING_DESK_DRAFT_WIDGET_KEY] = {
        "html": "<p>Frissen gépelt export mondat.</p>"
    }
    monkeypatch.setattr(st, "session_state", session)
    blob, name = writing_desk_ui.writing_desk_docx_export_payload()
    assert blob is not None
    assert name == "Textus_Jn_3_16_vazlat.docx"
    doc = Document(io.BytesIO(blob))
    text = _all_text(doc)
    assert "Frissen gépelt export mondat." in text
    assert "régi durable" not in text
    assert session[WRITING_DESK_KEY]["draft"]["content"] == (
        "<p>Frissen gépelt export mondat.</p>"
    )


def test_export_flush_respects_resync_pending_guard(monkeypatch):
    import streamlit as st

    import writing_desk_ui
    from writing_desk_ui import (
        WRITING_DESK_DRAFT_RESYNC_FLAG,
        WRITING_DESK_DRAFT_WIDGET_KEY,
    )

    session: dict = {"last_igehely": "Jn 3,16"}
    set_writing_desk_draft(session, "<p>új projekt</p>")
    session[WRITING_DESK_DRAFT_WIDGET_KEY] = {"html": "<p>stale widget</p>"}
    session[WRITING_DESK_DRAFT_RESYNC_FLAG] = True
    monkeypatch.setattr(st, "session_state", session)
    blob, _name = writing_desk_ui.writing_desk_docx_export_payload()
    doc = Document(io.BytesIO(blob))
    text = _all_text(doc)
    assert "új projekt" in text
    assert "stale widget" not in text


def test_empty_export_ui_disables_download(monkeypatch):
    import streamlit as st

    import writing_desk_ui

    session: dict = {}
    monkeypatch.setattr(writing_desk_ui, "_render_scripture_block", lambda: None)
    from test_writing_desk_ui import _patch_streamlit_shell

    calls = _patch_streamlit_shell(monkeypatch, st, session)
    writing_desk_ui.render_writing_desk_shell()
    assert calls["download_button"]
    assert calls["download_button"][0]["disabled"] is True
    assert "Nincs exportálható szöveg." in calls["caption"]
