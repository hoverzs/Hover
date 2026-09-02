"""Íróasztal draft → DOCX. Csak a 4B restricted HTML formázását őrzi.

Streamlit-független; a hívó flusholja a CCv2 widget-state-et export előtt.
"""

from __future__ import annotations

import io
import re
from html.parser import HTMLParser
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from writing_desk_data import (
    draft_has_visible_content,
    sanitize_draft_html,
)

_TEXT_ALIGN_RE = re.compile(
    r"text-align\s*:\s*(left|center|right|justify)\b",
    re.IGNORECASE,
)
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _paragraph_align(attrs: list[tuple[str, str | None]]) -> WD_ALIGN_PARAGRAPH | None:
    for raw_name, raw_value in attrs:
        name = (raw_name or "").lower()
        value = str(raw_value or "")
        if name == "style":
            match = _TEXT_ALIGN_RE.search(value)
            if match:
                return _ALIGN_MAP.get(match.group(1).lower())
        if name == "align":
            return _ALIGN_MAP.get(value.strip().lower())
    return None


_UNSAFE_FILENAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_NON_SLUG_RE = re.compile(r"[^\w]+", re.UNICODE)
_MULTIUSCORE_RE = re.compile(r"_+")


def writing_desk_docx_filename(reference: str | None) -> str:
    """Filesystem-safe Word fájlnév. A draft tartalmát nem használja."""
    raw = str(reference or "").strip()
    if not raw:
        return "Textus_vazlat.docx"
    slug = raw.replace(",", " ").replace(":", " ").replace(".", " ")
    slug = _UNSAFE_FILENAME_RE.sub(" ", slug)
    slug = _NON_SLUG_RE.sub("_", slug)
    slug = _MULTIUSCORE_RE.sub("_", slug).strip("_")
    if not slug:
        return "Textus_vazlat.docx"
    if len(slug) > 60:
        slug = slug[:60].rstrip("_")
    return f"Textus_{slug}_vazlat.docx"


def _clear_default_paragraph(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:p"):
            body.remove(child)


def _set_run_style(run, *, bold: bool, italic: bool, underline: bool) -> None:
    run.bold = True if bold else None
    run.italic = True if italic else None
    run.underline = True if underline else None
    run.font.name = "Calibri"
    run.font.size = Pt(11)


class _DraftHtmlToDocx(HTMLParser):
    def __init__(self, document: Document) -> None:
        super().__init__(convert_charrefs=True)
        self.doc = document
        self._para: Paragraph | None = None
        self._bold = 0
        self._italic = 0
        self._underline = 0
        self._list_stack: list[str] = []
        self._suppress_p = False

    def _list_style(self) -> str:
        kind = self._list_stack[-1] if self._list_stack else ""
        return "List Number" if kind == "ol" else "List Bullet"

    def _add_paragraph(
        self,
        *,
        list_item: bool = False,
        align: WD_ALIGN_PARAGRAPH | None = None,
    ) -> Paragraph:
        if list_item and self._list_stack:
            para = self.doc.add_paragraph(style=self._list_style())
            depth = max(len(self._list_stack) - 1, 0)
            if depth:
                para.paragraph_format.left_indent = Inches(0.25 * depth)
        else:
            para = self.doc.add_paragraph()
        if align is not None:
            para.alignment = align
        self._para = para
        return para

    def _ensure_para(self) -> Paragraph:
        if self._para is None:
            return self._add_paragraph(list_item=bool(self._list_stack))
        return self._para

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        name = (tag or "").lower()
        if name == "br":
            para = self._ensure_para()
            run = para.add_run()
            run.add_break()
            return
        if name in {"strong", "b"}:
            self._bold += 1
            return
        if name in {"em", "i"}:
            self._italic += 1
            return
        if name == "u":
            self._underline += 1
            return
        if name in {"ul", "ol"}:
            self._list_stack.append(name)
            self._para = None
            return
        if name == "li":
            self._add_paragraph(list_item=True, align=_paragraph_align(attrs))
            self._suppress_p = True
            return
        if name == "p":
            align = _paragraph_align(attrs)
            if self._suppress_p:
                self._suppress_p = False
                if self._para is None:
                    self._add_paragraph(list_item=bool(self._list_stack), align=align)
                elif align is not None:
                    self._para.alignment = align
                return
            self._add_paragraph(list_item=bool(self._list_stack), align=align)

    def handle_endtag(self, tag: str) -> None:
        name = (tag or "").lower()
        if name in {"strong", "b"}:
            self._bold = max(self._bold - 1, 0)
        elif name in {"em", "i"}:
            self._italic = max(self._italic - 1, 0)
        elif name == "u":
            self._underline = max(self._underline - 1, 0)
        elif name in {"ul", "ol"}:
            if self._list_stack:
                self._list_stack.pop()
            self._para = None
        elif name == "li":
            self._suppress_p = False
            self._para = None
        elif name == "p":
            if not self._list_stack:
                self._para = None

    def handle_data(self, data: str) -> None:
        if not data:
            return
        para = self._ensure_para()
        run = para.add_run(data)
        _set_run_style(
            run,
            bold=self._bold > 0,
            italic=self._italic > 0,
            underline=self._underline > 0,
        )


def build_writing_desk_docx_bytes(raw_html: Any) -> bytes | None:
    """Sanitize + 4B HTML → DOCX bytes. Üres / csak whitespace draft → None."""
    sanitized = sanitize_draft_html(raw_html)
    if not draft_has_visible_content(sanitized):
        return None
    document = Document()
    _clear_default_paragraph(document)
    parser = _DraftHtmlToDocx(document)
    try:
        parser.feed(sanitized)
        parser.close()
    except Exception:
        return None
    if not any((p.text or "").strip() for p in document.paragraphs):
        return None
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


__all__ = [
    "build_writing_desk_docx_bytes",
    "writing_desk_docx_filename",
]
